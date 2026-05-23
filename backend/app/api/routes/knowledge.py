import os
import re
from uuid import UUID

import pdfplumber
from docx import Document
from fastapi import APIRouter, Depends, HTTPException
from openpyxl import load_workbook
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.database import get_db
from app.core.deps import get_current_user
from app.models import File as FileModel
from app.models import KbChunk, KbDocument, User
from app.services.ai.client import embed_texts
from app.services.ai.prompts import SYSTEM_NO_HOSUN_BIAS

router = APIRouter(prefix="/knowledge", tags=["knowledge"])


class IngestBody(BaseModel):
    file_id: UUID
    title: str | None = None
    scope: str | None = "internal"


class QueryBody(BaseModel):
    question: str
    scope: str | None = None
    top_k: int = 5


def _read_text_from_file(path: str, mime: str | None) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                parts.append(t)
        return "\n".join(parts)
    if ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    if ext in (".xlsx", ".xls"):
        wb = load_workbook(path, read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                lines.append(" ".join(str(c) for c in row if c is not None))
        return "\n".join(lines)
    with open(path, encoding="utf-8", errors="ignore") as f:
        return f.read()


def _chunk_text(s: str, size: int = 1200) -> list[str]:
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return []
    return [s[i : i + size] for i in range(0, len(s), size)]


@router.post("/ingest")
def ingest(
    body: IngestBody,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
) -> dict:
    f = db.query(FileModel).filter(FileModel.id == body.file_id).first()
    if not f:
        raise HTTPException(status_code=404, detail="File not found")
    base = get_settings().UPLOAD_DIR
    abs_path = os.path.join(base, f.storage_key)
    if not os.path.isfile(abs_path):
        raise HTTPException(status_code=400, detail="File missing on disk")
    text_content = _read_text_from_file(abs_path, f.mime)
    chunks = _chunk_text(text_content)
    if not chunks:
        raise HTTPException(status_code=400, detail="No extractable text")
    doc = KbDocument(
        file_id=f.id,
        title=body.title or f.original_filename,
        source_type="upload",
        parse_status="parsed",
        embedding_model="text-embedding-3-small",
        scope=body.scope,
        created_by_id=user.id,
        updated_by_id=user.id,
    )
    db.add(doc)
    db.flush()
    vectors = embed_texts(chunks)
    for i, piece in enumerate(chunks):
        vec = vectors[i]
        ch = KbChunk(
            document_id=doc.id,
            chunk_index=i,
            content=piece,
            embedding=vec,
            created_by_id=user.id,
            updated_by_id=user.id,
        )
        db.add(ch)
    db.commit()
    return {"document_id": str(doc.id), "chunks": len(chunks)}


@router.post("/query")
def query_kb(
    body: QueryBody,
    db: Session = Depends(get_db),
    _: User = Depends(get_current_user),
) -> dict:
    if not body.question.strip():
        raise HTTPException(status_code=400, detail="Empty question")
    qv = embed_texts([body.question])[0]
    vec_literal = "[" + ",".join(str(float(x)) for x in qv) + "]"
    if body.scope:
        sql = text(
            """
            SELECT c.id, c.content, c.document_id
            FROM kb_chunks c
            JOIN kb_documents d ON d.id = c.document_id
            WHERE c.embedding IS NOT NULL AND d.scope = :scope
            ORDER BY c.embedding <=> CAST(:vec AS vector)
            LIMIT :k
            """
        )
        rows = db.execute(sql, {"vec": vec_literal, "k": body.top_k, "scope": body.scope}).mappings().all()
    else:
        sql = text(
            """
            SELECT c.id, c.content, c.document_id
            FROM kb_chunks c
            JOIN kb_documents d ON d.id = c.document_id
            WHERE c.embedding IS NOT NULL
            ORDER BY c.embedding <=> CAST(:vec AS vector)
            LIMIT :k
            """
        )
        rows = db.execute(sql, {"vec": vec_literal, "k": body.top_k}).mappings().all()
    cited = [
        {"chunk_id": str(r["id"]), "document_id": str(r["document_id"]), "excerpt": (r["content"] or "")[:500]}
        for r in rows
    ]
    from app.services.ai import client as ai_client

    context = "\n\n".join((r["content"] or "") for r in rows)
    msgs = [
        {"role": "system", "content": SYSTEM_NO_HOSUN_BIAS + " Answer using the context; cite chunk ids in brackets."},
        {"role": "user", "content": f"Question: {body.question}\n\nContext:\n{context}"},
    ]
    answer = ai_client.chat_completion(msgs)
    return {"answer": answer, "citations": cited}
